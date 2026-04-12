from __future__ import annotations

import os
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"c:\Users\user\Downloads\Najahni-symfony-main (3)\Najahni-symfony-main")
OUTPUT_DIR = ROOT / "docs"
OUTPUT_FILE = OUTPUT_DIR / "Najahni_Module_Investissement_Dossier_GL_S8_S10.docx"


@dataclass
class ProductStory:
    story_id: str
    priority: int
    role: str
    statement: str
    acceptance: list[str]
    points: int
    sprint: int
    status: str = "Termine"


def set_cell_text(cell, text: str, font_name: str = "Times New Roman", size: int = 9, bold: bool = False):
    cell.text = ""
    paragraphs = text.split("\n") if text else [""]
    for index, value in enumerate(paragraphs):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        run = paragraph.add_run(value)
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        run.font.size = Pt(size)
        run.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_cm: float):
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width_cm / 2.54 * 1440)))


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    default_text = OxmlElement("w:t")
    default_text.text = "Mettez a jour la table des matieres dans Word si necessaire."
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(default_text)
    run._r.append(fld_char_end)


def configure_styles(document: Document):
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(11)

    for style_name, size, color, bold in [
        ("Title", 24, RGBColor(18, 52, 86), True),
        ("Heading 1", 16, RGBColor(17, 69, 126), True),
        ("Heading 2", 13, RGBColor(31, 78, 121), True),
        ("Heading 3", 11, RGBColor(54, 95, 145), True),
    ]:
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = bold

    if "CodeBlock" not in document.styles:
        code_style = document.styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        code_style.base_style = document.styles["Normal"]
        code_style.font.name = "Consolas"
        code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        code_style.font.size = Pt(9)


def configure_page(document: Document):
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Page ")
    footer_run.font.name = "Calibri"
    footer_run.font.size = Pt(9)
    add_page_field(footer)


def add_title_page(document: Document):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(70)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("Najahni")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(17, 69, 126)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Module Investissement")
    run.font.name = "Calibri"
    run.font.size = Pt(21)
    run.bold = True

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Dossier Genie Logiciel - Semaines 8 a 10")
    run.font.name = "Calibri"
    run.font.size = Pt(17)
    run.bold = True

    for line in [
        "Application concrete des concepts Scrum, MVC et architecture en couches",
        "Projet de reference: plateforme entrepreneuriale tunisienne Najahni",
        "Base technique analysee: depot Najahni-symfony-main, module Investissement Symfony 6+/7",
        "Etudiant: 3eme annee d'ingenierie informatique - Esprit, Tunisie",
        "Date du livrable: 12 avril 2026",
    ]:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(line)
        run.font.size = Pt(12)

    box = document.add_paragraph()
    box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    box.paragraph_format.space_before = Pt(26)
    run = box.add_run(
        "Ce dossier montre que les concepts theoriques du cours ne sont pas restes abstraits: ils ont ete mobilises pour concevoir, organiser, implementer et finaliser un module professionnel couvrant tout le cycle de vie d'un investissement."
    )
    run.italic = True
    run.font.size = Pt(12)

    document.add_page_break()


def add_paragraph(document: Document, text: str, italic: bool = False, bold: bool = False, style: str | None = None):
    paragraph = document.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    run.italic = italic
    run.bold = bold
    return paragraph


def add_bullets(document: Document, items: list[str]):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.add_run(item)


def add_numbered(document: Document, items: list[str]):
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.add_run(item)


def add_landscape_section(document: Document):
    section = document.add_section(WD_SECTION_START.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    return section


def add_portrait_section(document: Document):
    section = document.add_section(WD_SECTION_START.NEW_PAGE)
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width, section.page_height = Inches(8.27), Inches(11.69)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    return section


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths_cm: list[float], font_size: int = 9):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(hdr[idx], header, font_name="Calibri", size=font_size, bold=True)
        shade_cell(hdr[idx], "D9EAF7")
        set_cell_width(hdr[idx], widths_cm[idx])

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=font_size)
            set_cell_width(cells[idx], widths_cm[idx])
    return table


def build_product_backlog() -> list[ProductStory]:
    return [
        ProductStory(
            "US-INV-01", 1, "Entrepreneur",
            "En tant qu'entrepreneur, je veux publier une opportunite d'investissement rattachee a l'un de mes projets afin de rechercher un financement cible dans un cadre structure.",
            [
                "L'utilisateur entrepreneur ne peut selectionner que ses propres projets.",
                "Le formulaire impose un montant cible >= 100 DT, une deadline future et une description significative.",
                "L'opportunite enregistree est creee avec le statut OPEN et apparait dans la liste publique des opportunites.",
            ],
            5, 1,
        ),
        ProductStory(
            "US-INV-02", 1, "Investisseur",
            "En tant qu'investisseur, je veux consulter les opportunites d'investissement avec recherche et tri AJAX afin d'identifier rapidement les projets les plus pertinents.",
            [
                "La liste affiche le titre du projet, la progression du financement, le montant cible et la deadline.",
                "Les filtres par texte et tri se declenchent sans rechargement complet de la page.",
                "Le compteur et le flux d'activite sont actualises avec les resultats filtres.",
            ],
            5, 1,
        ),
        ProductStory(
            "US-INV-03", 2, "Investisseur",
            "En tant qu'investisseur, je veux consulter la fiche detaillee d'une opportunite afin d'analyser son avancement, son contexte et les offres existantes.",
            [
                "La page montre la description complete, le pourcentage finance et la date limite.",
                "Les offres existantes sont visibles avec leur statut.",
                "Si l'utilisateur est entrepreneur proprietaire, il voit aussi les actions de gestion des offres.",
            ],
            3, 1,
        ),
        ProductStory(
            "US-INV-04", 1, "Investisseur",
            "En tant qu'investisseur, je veux lancer une analyse de risque IA sur une opportunite afin d'evaluer le niveau de risque avant de proposer des fonds.",
            [
                "Le calcul combine montant cible, duree restante et indicateurs economiques externes.",
                "Le resultat renvoie un score sur 100, un niveau de risque, une couleur et une recommandation.",
                "Le score calcule peut etre persiste sur l'opportunite pour reutilisation dans le module.",
            ],
            5, 1,
        ),
        ProductStory(
            "US-INV-05", 1, "Investisseur",
            "En tant qu'investisseur, je veux soumettre une offre d'investissement avec un gate de risque afin d'eviter un engagement non eclaire.",
            [
                "Le systeme bloque les auto-investissements sur son propre projet.",
                "Une offre en doublon sur la meme opportunite est refusee.",
                "Si le risque est modere ou eleve, une etape de confirmation explicite est demandee avant soumission.",
            ],
            5, 1,
        ),
        ProductStory(
            "US-INV-06", 2, "Entrepreneur",
            "En tant qu'entrepreneur, je veux visualiser les offres recues sur mes opportunites afin de comparer les propositions avant de prendre une decision.",
            [
                "Les offres sont visibles avec l'identite de l'investisseur, le montant, le statut et la date.",
                "Les offres en attente sont clairement distinguees des offres acceptees ou refusees.",
                "L'entrepreneur proprietaire est le seul a voir les commandes d'acceptation et de rejet.",
            ],
            3, 1,
        ),
        ProductStory(
            "US-INV-07", 1, "Entrepreneur",
            "En tant qu'entrepreneur, je veux accepter ou rejeter une offre afin de faire progresser uniquement les propositions retenues vers la contractualisation.",
            [
                "Une offre en attente peut etre acceptee ou refusee avec protection CSRF.",
                "Une notification est envoyee a l'investisseur apres la decision.",
                "Une opportunite peut passer a l'etat FUNDED si le financement cible est atteint.",
            ],
            3, 1,
        ),
        ProductStory(
            "US-INV-08", 1, "Investisseur",
            "En tant qu'investisseur, je veux qu'un espace contrat soit cree automatiquement apres acceptation de mon offre afin de formaliser la relation avec l'entrepreneur.",
            [
                "L'ouverture du contrat est reservee aux parties liees a l'offre acceptee.",
                "Le contrat est cree a la demande avec un titre, des termes par defaut et une empreinte SHA-256 initiale.",
                "Le statut initial du contrat est NEGOTIATING.",
            ],
            5, 2,
        ),
        ProductStory(
            "US-INV-09", 1, "Entrepreneur",
            "En tant qu'entrepreneur, je veux negocier les termes du contrat afin d'adapter l'accord aux besoins reels du projet et de l'investisseur.",
            [
                "Le titre, les termes, la consideration, les jalons textuels et l'equity peuvent etre modifies.",
                "Toute modification significative regenere le digest du contrat.",
                "Les signatures precedentes sont invalidees si les termes changent apres signature.",
            ],
            5, 2,
        ),
        ProductStory(
            "US-INV-10", 2, "Investisseur",
            "En tant qu'investisseur, je veux echanger des messages de negociation en temps reel dans le contrat afin d'eviter les echanges disperses hors plateforme.",
            [
                "Les messages sont charges chronologiquement et peuvent etre recuperes incrementalement.",
                "L'envoi d'un message met a jour l'activite du contrat et notifie l'autre partie.",
                "Les messages systeme sont differencies des messages manuels.",
            ],
            3, 2,
        ),
        ProductStory(
            "US-INV-11", 1, "Investisseur",
            "En tant qu'investisseur, je veux signer numeriquement le contrat avec une signature dessinee a la main et une preuve SHA-256 afin d'obtenir une trace forte de mon engagement.",
            [
                "La signature exige un nom complet et un dessin sur canvas encode en PNG base64.",
                "Le service calcule une empreinte SHA-256 a partir du digest du contrat et des metadonnees de signature.",
                "Quand les deux parties ont signe, le contrat passe a l'etat SIGNED et les vues PDF/print sont disponibles.",
            ],
            8, 2,
        ),
        ProductStory(
            "US-INV-12", 2, "Entrepreneur",
            "En tant qu'entrepreneur, je veux definir de 2 a 4 jalons de financement dans le contrat afin de lier les paiements a des livrables concrets.",
            [
                "Chaque jalon a un libelle, un pourcentage, un montant et une position.",
                "La somme des pourcentages doit etre exactement egale a 100%.",
                "Les jalons deviennent non modifiables des qu'un paiement partiel a deja ete libere.",
            ],
            5, 2,
        ),
        ProductStory(
            "US-INV-13", 2, "Entrepreneur",
            "En tant qu'entrepreneur, je veux marquer un jalon comme termine puis attendre la confirmation investisseur afin de suivre l'avancement du projet de facon controlee.",
            [
                "Seul l'entrepreneur peut passer un jalon de PENDING a COMPLETED.",
                "Seul l'investisseur peut le confirmer pour le faire passer a CONFIRMED.",
                "Chaque changement cree un message systeme et une notification metier.",
            ],
            3, 2,
        ),
        ProductStory(
            "US-INV-14", 3, "Investisseur",
            "En tant qu'investisseur, je veux verifier l'authenticite documentaire d'un contrat signe a partir d'une preuve courte exploitable par QR ou digest afin de confirmer que la version consultee est bien la bonne.",
            [
                "Le document contractuel affiche au minimum une empreinte SHA-256 unique et stable pour la version signee.",
                "La preuve est exploitable dans une logique de verification documentaire externe.",
                "Le mecanisme de verification reste coherent avec la generation du PDF signe.",
            ],
            2, 2,
        ),
        ProductStory(
            "US-INV-15", 1, "Investisseur",
            "En tant qu'investisseur, je veux payer un investissement accepte via Stripe, soit en une seule fois soit jalon par jalon, afin de financer le projet de maniere securisee.",
            [
                "Le paiement global n'est autorise que si le contrat est integralement signe et sans jalons definis.",
                "Le paiement par jalon n'est autorise qu'apres confirmation du jalon.",
                "Les identifiants Stripe et les dates de paiement sont conserves pour la tracabilite.",
            ],
            8, 3,
        ),
        ProductStory(
            "US-INV-16", 1, "Investisseur",
            "En tant qu'investisseur, je veux consulter mon portfolio d'investissement afin de suivre mes montants engages, mes jalons et la valeur projetee de mes participations.",
            [
                "Le portfolio ne liste que les offres effectivement payees.",
                "Chaque carte affiche le montant, la participation, la progression des jalons et une timeline de financement.",
                "Une synthese globale calcule le total investi et les transactions confirmees.",
            ],
            5, 3,
        ),
        ProductStory(
            "US-INV-17", 2, "Investisseur",
            "En tant qu'investisseur, je veux consulter un dashboard economique afin de contextualiser mes decisions d'investissement par secteur et par climat macroeconomique.",
            [
                "Le dashboard presente la repartition par risque, secteur et echeance courte.",
                "Un verdict synthese est construit a partir des indicateurs economiques externes.",
                "Les donnees economiques peuvent aussi etre exposees en JSON pour des composants dynamiques.",
            ],
            3, 3,
        ),
        ProductStory(
            "US-INV-18", 1, "Investisseur",
            "En tant qu'investisseur, je veux enregistrer mon profil puis recevoir un matching automatique avec les opportunites ouvertes afin de gagner du temps dans la prospection.",
            [
                "Le profil sauvegarde les secteurs preferes, la tolerance au risque, le budget et l'horizon.",
                "Le score de compatibilite est calcule sur 100 avec explication textuelle.",
                "Les recommandations sont restituées dans une interface dynamique sans pagination lourde.",
            ],
            5, 3,
        ),
        ProductStory(
            "US-INV-19", 2, "Investisseur",
            "En tant qu'investisseur, je veux interroger un chatbot IA contextuel afin d'obtenir des recommandations adaptees a mon opportunite, mon contrat ou mon profil.",
            [
                "Le chatbot refuse les messages vides et limite la longueur des prompts.",
                "Le contexte de l'opportunite ou du contrat est injecte dans la requete IA.",
                "Une reponse de secours coherente existe en cas d'indisponibilite du service IA.",
            ],
            5, 3,
        ),
        ProductStory(
            "US-INV-20", 2, "Administrateur",
            "En tant qu'administrateur, je veux consulter des listes filtrees d'opportunites et d'offres afin de superviser l'activite du module et de detecter les anomalies de cycle de vie.",
            [
                "Les repositories exposent des requetes filtrees par statut et recherche texte.",
                "Les indicateurs agreges comptent les opportunites, offres et statuts principaux.",
                "La supervision permet de reperer les contrats finances, signes ou encore en negociation.",
            ],
            3, 3,
        ),
        ProductStory(
            "US-INV-21", 3, "Administrateur",
            "En tant qu'administrateur, je veux disposer d'une vue d'audit sur les traces de paiement, de signature et de progression des contrats afin de garantir la conformite operationnelle du module.",
            [
                "Les identifiants de paiement, dates de signature et statuts de contrat sont conserves dans les entites metier.",
                "Les messages systeme offrent un journal fonctionnel de la negociation et des jalons.",
                "Les donnees restent consultables via les repositories et les vues metier du module.",
            ],
            2, 3,
        ),
    ]


def backlog_rows(backlog: list[ProductStory]) -> list[list[str]]:
    rows = []
    for story in backlog:
        acceptance = "\n".join(f"- {item}" for item in story.acceptance)
        rows.append([
            story.story_id,
            str(story.priority),
            story.statement,
            acceptance,
            str(story.points),
            str(story.sprint),
            story.status,
        ])
    return rows


def sprint2_tasks() -> list[tuple[str, str, list[tuple[str, int]]]]:
    return [
        (
            "US-INV-08",
            "Creation automatique de l'espace contrat",
            [
                ("Verifier l'acces a partir d'une offre acceptee et resoudre les parties investisseur/entrepreneur", 3),
                ("Coder la creation lazy du contrat avec titre, termes par defaut et digest initial", 5),
                ("Construire la vue contractuelle initiale et ses indicateurs d'etat", 4),
                ("Tester l'ouverture par les deux roles et les refus d'acces", 2),
            ],
        ),
        (
            "US-INV-09",
            "Negociation des termes contractuels",
            [
                ("Concevoir le formulaire de mise a jour des clauses et de l'equity", 3),
                ("Valider les contraintes fonctionnelles sur le titre, les termes et le pourcentage", 4),
                ("Regenerer l'empreinte SHA-256 et invalider les signatures en cas de changement", 4),
                ("Envoyer la notification a l'autre partie et tracer un message systeme", 2),
            ],
        ),
        (
            "US-INV-10",
            "Messagerie de negociation en temps reel",
            [
                ("Creer l'entite de message contractuel et le repository chronologique", 3),
                ("Implementer les endpoints GET et POST de messages en JSON", 4),
                ("Ajouter le polling cote interface et la mise a jour incrémentale du fil", 5),
                ("Gerer les notifications et la securisation CSRF", 2),
            ],
        ),
        (
            "US-INV-11",
            "Signature numerique SHA-256 avec canvas",
            [
                ("Ajouter le panneau de signature avec canvas, nom complet et export PNG base64", 5),
                ("Concevoir le service de calcul de digest et de hash de preuve de signature", 5),
                ("Persister les signatures, dates et images pour les deux parties", 4),
                ("Debloquer les vues PDF et impression apres double signature", 3),
                ("Tester les cas invalides: nom court, canvas vide, acces non autorise", 2),
            ],
        ),
        (
            "US-INV-12",
            "Definition des jalons de financement",
            [
                ("Modeliser ContractMilestone et son repository", 3),
                ("Construire le formulaire multi-jalons avec somme a 100%", 4),
                ("Calculer les montants derives a partir du pourcentage et du montant de l'offre", 3),
                ("Interdire l'edition des jalons apres liberation d'un paiement", 2),
            ],
        ),
        (
            "US-INV-13",
            "Workflow jalon termine, confirme, libere",
            [
                ("Coder les transitions PENDING -> COMPLETED -> CONFIRMED -> RELEASED", 4),
                ("Appliquer les controles de role entrepreneur/investisseur sur chaque etape", 3),
                ("Tracer les evenements dans la conversation contractuelle", 2),
                ("Notifier les parties selon l'etat du jalon", 2),
            ],
        ),
        (
            "US-INV-14",
            "Preuve d'authenticite documentaire du contrat",
            [
                ("Afficher le digest SHA-256 dans le document PDF et l'ecran contrat", 2),
                ("Structurer une preuve courte reutilisable pour la verification documentaire", 2),
                ("Verifier la coherence entre version signee, PDF genere et preuve exposee", 2),
            ],
        ),
    ]


def add_section_title(document: Document, title: str, level: int = 1):
    document.add_heading(title, level=level)


def add_intro_sections(document: Document):
    add_section_title(document, "Table des matieres", 1)
    toc_paragraph = document.add_paragraph()
    add_toc(toc_paragraph)
    document.add_page_break()

    add_section_title(document, "1. Introduction et positionnement du dossier", 1)
    add_paragraph(
        document,
        "Ce dossier a ete concu comme un livrable academique complet pour le cours Genie Logiciel & Atelier GL. Son objectif n'est pas seulement de decrire le module Investissement de Najahni, mais de demontrer que les concepts vus en cours ont ete compris puis appliques de maniere intentionnelle dans un projet Symfony de niveau professionnel."
    )
    add_paragraph(
        document,
        "L'analyse est basee sur le depot transmis, en particulier sur les entites InvestmentOpportunity, InvestmentOffer, InvestmentContract, InvestmentContractMessage, ContractMilestone et InvestorProfile, sur les controleurs InvestmentController, InvestmentContractController et InvestmentAdvancedController, sur les repositories Doctrine associes, sur les services metier du sous-dossier src/Service/Investment et sur les vues Twig du dossier templates/front/investment."
    )
    add_paragraph(
        document,
        "Le dossier est organise en deux grandes parties. La premiere traite la dimension Scrum: vision produit, Product Backlog, Sprint Backlog et Definition of Done. La seconde traite la dimension Architecture & Conception: architecture logique versus physique, articulation 3 couches et 5 couches, patron MVC, circulation inter-couches et justification des choix techniques."
    )

    add_section_title(document, "2. Perimetre fonctionnel du module Investissement", 1)
    add_paragraph(
        document,
        "Le module Investissement couvre un cycle de vie complet: publication d'opportunites, consultation avec filtres AJAX, evaluation du risque via un moteur economique, soumission d'offres avec gate de risque, acceptation ou rejet par l'entrepreneur, creation et negociation du contrat, messagerie embarquee, signature numerique, gestion de jalons de paiement, paiement Stripe, suivi des contrats, portfolio investisseur, matching et assistant IA contextuel."
    )
    add_bullets(
        document,
        [
            "Public cible: investisseurs, entrepreneurs et administrateurs.",
            "Socle technique principal: Symfony 7.2, Doctrine ORM, Twig, Bootstrap 5, JavaScript vanilla, MySQL.",
            "Services externes mobilises: Stripe, Hugging Face router, World Bank API, Open Exchange Rates, Dompdf.",
            "Preuve d'integrite contractuelle observee dans le code: empreinte SHA-256 et document PDF signe.",
        ],
    )


def add_scrum_part(document: Document, backlog: list[ProductStory]):
    add_section_title(document, "3. Partie 1 - Application de Scrum au module Investissement", 1)

    add_section_title(document, "3.1 Vision produit et logique de decoupage en sprints", 2)
    add_paragraph(
        document,
        "Du point de vue Scrum, le module Investissement peut etre vu comme un increment produit progressif construit en trois sprints. Le Sprint 1 vise la mise en marche du flux coeur: creation d'opportunites, consultation, analyse du risque et soumission d'offres. Le Sprint 2 consolide la contractualisation: negotiation, messagerie, signature numerique et structuration des jalons. Le Sprint 3 finalise la valeur metier: paiements Stripe, portfolio, matching, dashboard et assistant IA."
    )
    add_paragraph(
        document,
        "Cette structuration est defendable pedagogiquement parce qu'elle respecte un principe d'incrementation de valeur. Chaque sprint apporte une fonctionnalite demonstrable, testable et utile, tout en limitant la dette d'integration."
    )

    add_section_title(document, "3.2 Product Backlog complet du module", 2)
    add_paragraph(
        document,
        "Le Product Backlog ci-dessous reformule le module sous forme de User Stories officielles. Toutes les stories sont marquees Terminees car le module a deja ete implemente dans le projet. Les priorites sont exprimees selon la convention du cours: 1 pour haute, 2 pour moyenne et 3 pour basse."
    )

    add_landscape_section(document)
    add_section_title(document, "Tableau 1 - Product Backlog", 2)
    add_table(
        document,
        ["ID", "Priorite", "User Story", "Criteres d'acceptation", "Points", "Sprint", "Statut"],
        backlog_rows(backlog),
        [2.0, 1.6, 8.3, 10.8, 1.5, 1.5, 1.8],
        font_size=8,
    )

    add_portrait_section(document)
    add_section_title(document, "3.3 Lecture analytique du backlog", 2)
    add_paragraph(
        document,
        "Le backlog montre une progression logique du besoin metier. Les stories a priorite 1 correspondent aux fonctions vitales sans lesquelles la proposition de valeur n'existe pas: publier une opportunite, soumettre une offre, accepter une offre, ouvrir un contrat, signer, payer et suivre le portfolio. Les stories de priorite 2 enrichissent le pilotage et l'aide a la decision: risk engine, dashboard, matching, supervision. Les stories de priorite 3 concernent surtout les mecanismes transverses de verification et d'audit, importants mais non bloquants pour la premiere mise en service."
    )

    add_section_title(document, "3.4 Sprint Backlog detaille du Sprint 2", 2)
    add_paragraph(
        document,
        "Le Sprint 2 est celui de la contractualisation. Il constitue le coeur du module, car il transforme une offre acceptee en accord executable. Le Sprint Backlog ci-dessous extrait toutes les stories affectees au Sprint 2 et les decompose en taches techniques concretes."
    )

    sprint_rows = []
    for story_id, title, tasks in sprint2_tasks():
        total = sum(hours for _, hours in tasks)
        sprint_rows.append([
            story_id,
            title,
            "\n".join(f"- {task}" for task, _ in tasks),
            "\n".join(f"{hours} h" for _, hours in tasks),
            f"{total} h",
        ])

    add_landscape_section(document)
    add_section_title(document, "Tableau 2 - Sprint Backlog Sprint 2", 2)
    add_table(
        document,
        ["User Story", "Intitule", "Taches techniques", "Estimation par tache", "Total"],
        sprint_rows,
        [2.0, 5.5, 11.3, 3.0, 2.0],
        font_size=8,
    )

    add_portrait_section(document)
    add_section_title(document, "3.5 Definition of Done adaptee a Symfony/PHP", 2)
    add_paragraph(
        document,
        "Pour qu'une User Story du module Investissement soit declaree Terminee, l'equipe doit satisfaire une Definition of Done stricte. Cette definition garantit que la story n'est pas seulement codee, mais aussi integree, verifiee et coherente avec l'architecture Symfony du projet."
    )
    add_numbered(
        document,
        [
            "Le code PHP, Twig et JavaScript necessaire a la User Story est ecrit, relu et compréhensible pour un autre membre de l'equipe.",
            "Les validations serveur sont presentes: controle des types, des bornes metier, des roles, des conditions d'acces et des tokens CSRF.",
            "Le code PHP passe une verification de syntaxe type `php -l` ou equivalent sur les fichiers modifies.",
            "Les routes Symfony sont declarees, resolues et verifiees manuellement depuis l'interface ou par URL directe.",
            "Les templates Twig associes sont rendus sans erreur et les blocs HTML conditionnels ont ete controles.",
            "Le schema de base de donnees est a jour si la story introduit une nouvelle entite, une nouvelle colonne ou une nouvelle relation Doctrine.",
            "Les tests fonctionnels manuels ont ete executes sur les cas nominaux et sur les cas d'erreur evidents.",
            "Le cache Symfony est vide ou regenere apres les modifications structurelles pour eviter un faux positif d'integration.",
            "Les messages metier visibles par l'utilisateur, les notifications et les retours JSON sont coherents et non ambigus.",
            "La story peut etre montree au Product Owner sous forme d'une demonstration complete de bout en bout.",
        ],
    )

    add_section_title(document, "3.6 Pourquoi cette application de Scrum est credible", 2)
    add_paragraph(
        document,
        "Le backlog propose n'est pas artificiel. Il correspond reellement a la facon dont un module de ce type se construit. Les stories sont verticales, c'est-a-dire qu'elles livrent de la valeur fonctionnelle visible plutot qu'une simple liste de composants techniques. Les estimations en points restent raisonnables, les priorites distinguent le coeur du produit des enrichissements, et le Sprint 2 concentre bien le moment ou la complexite fonctionnelle augmente fortement avec la negotiation, la signature et les jalons."
    )


def add_architecture_part(document: Document):
    add_section_title(document, "4. Partie 2 - Architecture appliquee au module Investissement", 1)

    add_section_title(document, "4.1 Architecture logique versus architecture physique", 2)
    add_paragraph(
        document,
        "Dans le cours, l'architecture logique repond a la question 'quelles responsabilites le logiciel se repartit-il ?', alors que l'architecture physique repond a la question 'ou s'executent les composants et sur quelles ressources ?'. Le module Investissement illustre tres bien cette distinction. Logiquement, on distingue des vues Twig, des controleurs applicatifs, des services metier, des repositories Doctrine et des integrations externes. Physiquement, l'ensemble est deploye sur une application Symfony qui dialogue avec une base MySQL et plusieurs API distantes."
    )

    add_section_title(document, "4.1.1 Vue logique du module", 3)
    logical_rows = [
        ["Presentation", "templates/front/investment/*.html.twig + JavaScript vanilla", "Afficher les pages, capter les actions utilisateur, envoyer les requetes AJAX et presenter les resultats metier."],
        ["Logique applicative", "InvestmentController.php, InvestmentContractController.php, InvestmentAdvancedController.php", "Orchestrer les cas d'usage, securiser l'acces, appeler les services, choisir la vue ou la reponse JSON."],
        ["Metier", "EconomicRiskEngine.php, InvestmentMatchingService.php, ContractSignatureService.php, StripePaymentService.php, InvestmentChatbotService.php", "Porter les regles metier stables, calculer les scores, gerer la signature, piloter les integrations de paiement et d'IA."],
        ["Acces aux donnees", "InvestmentOpportunityRepository.php, InvestmentOfferRepository.php, InvestmentContractRepository.php, InvestmentContractMessageRepository.php, ContractMilestoneRepository.php", "Executer les requetes Doctrine et structurer la persistence sans polluer les controleurs."],
        ["Infrastructure", "Doctrine/MySQL, Stripe, Hugging Face, World Bank, Open Exchange Rates, Dompdf", "Fournir les ressources externes, la persistence et la production documentaire."],
    ]
    add_table(document, ["Bloc logique", "Elements concrets", "Role dans le module"], logical_rows, [3.2, 6.5, 7.8], font_size=9)

    add_section_title(document, "4.1.2 Vue physique du module", 3)
    physical_rows = [
        ["Navigateur web", "Pages Twig rendues par Symfony, composants Bootstrap, appels fetch AJAX vers les endpoints du module."],
        ["Serveur applicatif", "Application Symfony 7.2 executant les controleurs, services et templates du module Investissement."],
        ["Base de donnees", "MySQL via Doctrine ORM pour les entites InvestmentOpportunity, InvestmentOffer, InvestmentContract, InvestmentContractMessage, ContractMilestone et InvestorProfile."],
        ["Services externes", "Stripe pour le paiement, Hugging Face pour le chatbot, World Bank et Open Exchange Rates pour les donnees economiques, Dompdf pour la generation PDF."],
        ["Capacite documentaire additionnelle", "La plateforme Najahni dispose deja d'une integration QR sur un autre module; pour l'investissement, l'extrait de code confirme surtout l'usage du digest SHA-256 et du PDF signe comme preuve primaire."],
    ]
    add_table(document, ["Noeud physique", "Observation"], physical_rows, [4.0, 13.5], font_size=9)

    add_section_title(document, "4.2 Rappel du modele en 3 couches et son articulation avec le module", 2)
    add_paragraph(
        document,
        "Le modele en 3 couches du cours se retrouve naturellement dans Najahni. La couche presentation correspond aux vues Twig et au JavaScript embarque. La couche metier correspond a l'ensemble forme par les controleurs et surtout les services metier. La couche donnees correspond aux entites, repositories Doctrine et a MySQL. Le modele en 5 couches affine cette vision en separant plus proprement l'orchestration applicative, le coeur metier et l'infrastructure."
    )

    add_section_title(document, "4.3 Patron MVC implemente dans Symfony", 2)
    add_paragraph(
        document,
        "Le framework Symfony applique nativement le patron MVC. Le module Investissement le suit de facon lisible: les entites jouent le role de Modele, les templates Twig jouent le role de Vue et les controleurs HTTP jouent le role de Controleur. L'interet de ce patron est double: separer les responsabilites et rendre le module evolutif sans dupliquer la logique."
    )

    add_section_title(document, "4.3.1 Couche Modele", 3)
    model_rows = [
        ["InvestmentOpportunity.php", "Represente l'opportunite ouverte a l'investissement: montant cible, description, deadline, statut, score de risque, relation au projet et collection d'offres."],
        ["InvestmentOffer.php", "Represente la proposition de l'investisseur: montant propose, statut PENDING/ACCEPTED/REJECTED, indicateurs de paiement et drapeau de risque acknowledge."],
        ["InvestmentContract.php", "Represente le coeur contractuel du module: parties, clauses, equity, digest, signatures, statut et liste des jalons de financement."],
        ["InvestmentContractMessage.php", "Represente les echanges de negociation ou les messages systeme lies a un contrat."],
        ["ContractMilestone.php", "Represente un jalon de financement avec pourcentage, montant, position et etat de progression."],
        ["InvestorProfile.php", "Represente le profil d'investissement servant au matching: secteurs, budget, tolerance au risque et horizon temporel."],
    ]
    add_table(document, ["Fichier Modele", "Justification MVC"], model_rows, [5.2, 12.3], font_size=9)
    add_paragraph(
        document,
        "Ces entites incarnent bien la couche Modele parce qu'elles portent l'etat durable du domaine. Elles ne dessinent aucune interface et n'orchestrent pas la navigation. Elles representent le langage metier du module: opportunite, offre, contrat, message, jalon, profil investisseur."
    )

    add_section_title(document, "4.3.2 Couche Vue", 3)
    view_rows = [
        ["opportunities.html.twig", "Vue catalogue qui affiche la liste des opportunites, les filtres AJAX, le deal feed et les actions principales de navigation."],
        ["show.html.twig", "Vue detail d'une opportunite, incluant progression, liste des offres, formulaire d'investissement et modal de risk gate."],
        ["contract.html.twig", "Vue de negociation contractuelle avec clauses, signatures, messagerie, stepper de cycle de vie et actions de jalons."],
        ["portfolio.html.twig", "Vue de synthese investisseur presentant les cartes d'investissement, les timelines et les indicateurs consolides."],
        ["risk_analysis.html.twig", "Vue analytique qui illustre la restitution du score, des facteurs et du verdict IA dans une interface riche."],
        ["matching.html.twig", "Vue de profil + recommandations qui combine formulaire, fetch AJAX et cartes dynamiques de compatibilite."],
    ]
    add_table(document, ["Template Vue", "Role dans MVC"], view_rows, [4.8, 12.7], font_size=9)
    add_paragraph(
        document,
        "On voit clairement la responsabilite de la Vue dans les templates. Par exemple, opportunities.html.twig contient le champ de recherche, le select de tri et un script fetch qui appelle l'endpoint AJAX sans contenir lui-meme de requete Doctrine. De la meme facon, contract.html.twig met en scene l'etat du contrat, mais ne calcule pas l'empreinte ni la logique de signature."
    )

    add_section_title(document, "4.3.3 Couche Controleur", 3)
    controller_rows = [
        ["InvestmentController.php", "Orchestre les cas d'usage coeur: consultation des opportunites, detail, soumission d'offre, paiement, creation d'opportunite, portfolio et listes personnelles."],
        ["InvestmentContractController.php", "Pilote le sous-cycle contractuel: creation lazy du contrat, mise a jour des termes, signature, messagerie, jalons, PDF, impression et assistant de negociation."],
        ["InvestmentAdvancedController.php", "Porte les cas d'usage avances: dashboard economique, calcul de risque, matching, sauvegarde de profil et chatbot IA contextuel."],
    ]
    add_table(document, ["Controleur", "Role d'orchestration"], controller_rows, [4.6, 12.9], font_size=9)
    add_paragraph(
        document,
        "Le role du Controleur est lisible dans le code. Exemple concret: la route /investissement/opportunities/ajax declenche opportunitiesAjax(), qui recupere les parametres HTTP, interroge InvestmentOpportunityRepository::searchOpen(), enrichit les donnees avec EconomicApiService puis renvoie une reponse JSON exploitable par la Vue."
    )

    add_section_title(document, "4.4 Architecture en 5 couches appliquee au module", 2)
    add_paragraph(
        document,
        "Le modele en 5 couches du cours permet une lecture encore plus precise. Il montre que le module n'est pas seulement un empilement de fichiers Symfony, mais une architecture structuree dans laquelle chaque couche a une responsabilite nette."
    )

    add_section_title(document, "4.4.1 Couche IHM", 3)
    add_paragraph(
        document,
        "La couche IHM est composee des templates Twig du dossier templates/front/investment, du JavaScript vanilla embarque dans les pages et de Bootstrap 5 pour la mise en page et les composants visuels. Cette couche gere les formulaires, les listes, les animations, les modales, les timelines et les appels AJAX. Elle ne prend pas de decision metier durable; elle prepare les donnees d'entree et presente les resultats."
    )
    add_paragraph(
        document,
        "Exemple concret: opportunities.html.twig capte la saisie utilisateur, lance fetch vers l'endpoint app_invest_opportunities_ajax et reconstruit dynamiquement les cartes d'opportunites. matching.html.twig enregistre d'abord le profil investisseur, puis recharge les recommandations en JSON. contract.html.twig embarque un canvas de signature et une messagerie dynamique."
    )

    add_section_title(document, "4.4.2 Couche Logique Applicative", 3)
    add_paragraph(
        document,
        "La couche logique applicative est portee principalement par InvestmentController.php, InvestmentContractController.php et InvestmentAdvancedController.php. Elle transforme une intention HTTP en cas d'usage executable: verifier les droits, normaliser les parametres, appeler les services, persister les changements, choisir la reponse HTML ou JSON et declencher les notifications."
    )
    add_paragraph(
        document,
        "Exemple concret: lors d'un paiement global, InvestmentController::payOffer() verifie d'abord l'investisseur courant, le token CSRF, le statut de l'offre et la presence d'un contrat integralement signe. Ensuite seulement il appelle StripePaymentService, met a jour l'entite InvestmentOffer et declenche des notifications metier."
    )

    add_section_title(document, "4.4.3 Couche Metier", 3)
    add_paragraph(
        document,
        "La couche metier contient les regles stables du domaine et les calculs reutilisables. EconomicRiskEngine.php calcule un score de risque en combinant montant, duree et contexte economique. InvestmentMatchingService.php calcule un score de compatibilite a partir du secteur, du budget, du risque et de l'horizon. ContractSignatureService.php gere le digest contractuel et la preuve cryptographique de signature. StripePaymentService.php encapsule la logique de creation d'un PaymentIntent. InvestmentChatbotService.php construit les prompts systeme et gere l'appel a Hugging Face."
    )
    add_paragraph(
        document,
        "Exemple concret d'interaction: InvestmentAdvancedController::riskCompute() appelle EconomicApiService pour obtenir les indicateurs macroeconomiques, puis appelle EconomicRiskEngine::calculateFullRisk(). Le resultat est ensuite persiste dans l'entite InvestmentOpportunity avant d'etre renvoye a l'IHM."
    )

    add_section_title(document, "4.4.4 Couche Acces Donnees", 3)
    add_paragraph(
        document,
        "La couche acces donnees est portee par les repositories Doctrine. InvestmentOpportunityRepository expose searchOpen(), countOpenByRiskBracket() et countOpenBySector(). InvestmentOfferRepository expose findExistingOffer(), findPaidByInvestor() et findUnpaidByInvestor(). InvestmentContractRepository recentre la recherche de contrats par utilisateur. InvestmentContractMessageRepository apporte un acces chronologique ou incremental aux messages. ContractMilestoneRepository structure l'acces aux jalons et aux sommes deja liberees."
    )
    add_paragraph(
        document,
        "Exemple concret: searchOpen() centralise la logique de filtrage et de tri des opportunites. Ce choix evite de repliquer des conditions de requetage dans plusieurs controleurs et renforce la cohesion de la couche persistence."
    )

    add_section_title(document, "4.4.5 Couche Infrastructure", 3)
    add_paragraph(
        document,
        "La couche infrastructure correspond aux ressources physiques ou externes dont depend le module: MySQL via Doctrine, Stripe via stripe-php, Hugging Face via HttpClient, World Bank API et Open Exchange Rates pour l'economie, ainsi que Dompdf pour la production du document contractuel. Elle comprend egalement les mecanismes de transport HTTP et de serialisation necessaires aux appels externes."
    )
    add_paragraph(
        document,
        "Exemple concret: EconomicApiService::fetchAllEconomicData() appelle successivement le service de taux de change puis deux indicateurs World Bank. StripePaymentService::payAcceptedOffer() cree un PaymentIntent Stripe. InvestmentContractController::pdf() rend contract_pdf.html.twig puis genere un PDF via Dompdf."
    )

    add_section_title(document, "4.5 Scenarios de circulation entre les couches", 2)
    scenario_rows = [
        [
            "Consultation et filtrage des opportunites",
            "IHM: opportunities.html.twig envoie fetch AJAX -> Logique applicative: InvestmentController::opportunitiesAjax() -> Acces donnees: InvestmentOpportunityRepository::searchOpen() -> Metier: buildEcoBadge() et buildPitchLine() + EconomicApiService -> IHM: rendu dynamique des cartes."
        ],
        [
            "Analyse de risque d'une opportunite",
            "IHM: risk_analysis.html.twig appelle l'endpoint de calcul -> Logique applicative: InvestmentAdvancedController::riskCompute() -> Infrastructure: EconomicApiService interroge World Bank/Open Exchange -> Metier: EconomicRiskEngine calcule le score -> Acces donnees: persistence du score sur InvestmentOpportunity -> IHM: affichage du verdict."
        ],
        [
            "Contractualisation et signature",
            "IHM: contract.html.twig poste les termes et les signatures -> Logique applicative: InvestmentContractController -> Metier: ContractSignatureService genere digest et hash SHA-256 -> Acces donnees: InvestmentContract et InvestmentContractMessage sont persistes -> Infrastructure: Dompdf rend le PDF signe lorsque les deux signatures existent."
        ],
        [
            "Paiement par jalon",
            "IHM: action sur le jalon confirme -> Logique applicative: milestoneRelease() -> Metier: StripePaymentService cree un PaymentIntent -> Acces donnees: mise a jour de ContractMilestone, InvestmentOffer et InvestmentContract -> IHM: timeline et statut de financement mis a jour."
        ],
        [
            "Matching IA",
            "IHM: matching.html.twig sauvegarde le profil puis recharge les resultats -> Logique applicative: saveProfile() puis matchingResults() -> Acces donnees: InvestorProfileRepository et InvestmentOpportunityRepository -> Metier: InvestmentMatchingService::findMatches() -> IHM: cartes de compatibilite et liens vers l'analyse de risque."
        ],
    ]
    add_table(document, ["Scenario", "Flux inter-couches observe"], scenario_rows, [5.0, 12.5], font_size=9)

    add_section_title(document, "4.6 Justification des choix architecturaux", 2)

    add_section_title(document, "4.6.1 Pourquoi MVC avec Symfony", 3)
    add_paragraph(
        document,
        "Le choix de MVC avec Symfony est pertinent parce que le module presente a la fois une richesse d'IHM et une logique metier importante. En separant clairement Modeles, Vues et Controleurs, l'equipe peut faire evoluer l'interface Twig sans toucher a la persistence, ou faire evoluer la logique metier sans casser le rendu. Cette separation diminue le couplage, facilite les tests manuels et rend le code plus compréhensible pour une equipe etudiante travaillant en groupe."
    )

    add_section_title(document, "4.6.2 Pourquoi une couche metier distincte", 3)
    add_paragraph(
        document,
        "La separation des services metier hors des controleurs est une decision essentielle. Un controleur HTTP ne doit pas contenir toute la connaissance du domaine. En deplaçant le calcul du risque, le matching, la signature ou le paiement dans des services dedies, on augmente la cohesion des classes, on evite la duplication et on rend les regles reutilisables depuis plusieurs cas d'usage. C'est exactement l'un des gains recherches par l'architecture en couches: localiser la complexite au bon endroit."
    )

    add_section_title(document, "4.6.3 Pourquoi des repositories dedies", 3)
    add_paragraph(
        document,
        "Les repositories dedies structurent l'acces aux donnees et evitent l'eparpillement des requetes Doctrine dans les controleurs. Ce choix ameliore la maintenabilite, car une evolution de filtre, de tri ou d'agregation ne se fait qu'a un seul endroit. Il soutient aussi la lisibilite pedagogique du projet: un professeur peut identifier rapidement ou se trouvent les requetes de recherche, de comptage ou de supervision du module."
    )

    add_section_title(document, "4.6.4 Pourquoi une architecture en couches favorise la maintenabilite et la reutilisabilite", 3)
    add_paragraph(
        document,
        "L'architecture en couches favorise la maintenabilite parce qu'elle isole les impacts de changement. Modifier l'IHM n'impose pas de re-ecrire le moteur de risque. Changer l'API economique n'impose pas de re-dessiner les templates. Ajouter un nouveau canal de signature ou un autre mode de paiement resterait localise dans la couche infrastructure ou metier. Elle favorise egalement la reutilisabilite: EconomicRiskEngine, InvestmentMatchingService et ContractSignatureService peuvent etre mobilises par d'autres ecrans, par des commandes console ou par de futurs tests automatises. Enfin, elle soutient la comprehension collective, ce qui est central dans un projet de groupe gere en Scrum."
    )

    add_section_title(document, "4.7 Preuves techniques selectionnees dans le depot", 2)
    evidence_items = [
        "src/Controller/InvestmentController.php: publication d'opportunites, consultation AJAX, offres, portfolio, paiement Stripe global.",
        "src/Controller/InvestmentContractController.php: creation du contrat, negociation, signatures, messages, jalons, PDF et impression.",
        "src/Controller/InvestmentAdvancedController.php: dashboard economique, risque, matching, profil investisseur, chatbot.",
        "src/Service/Investment/EconomicRiskEngine.php: calcul deterministe du risque.",
        "src/Service/Investment/InvestmentMatchingService.php: score de compatibilite investisseur-opportunite.",
        "src/Service/Investment/ContractSignatureService.php: digest et preuves SHA-256.",
        "src/Service/Investment/StripePaymentService.php: encapsulation de Stripe.",
        "src/Service/Investment/InvestmentChatbotService.php et EconomicApiService.php: IA contextuelle et collecte des donnees macroeconomiques.",
        "templates/front/investment/opportunities.html.twig, show.html.twig, contract.html.twig, risk_analysis.html.twig, matching.html.twig, portfolio.html.twig: preuves UI du flux complet.",
    ]
    add_bullets(document, evidence_items)


def add_conclusion(document: Document):
    add_section_title(document, "5. Conclusion generale", 1)
    add_paragraph(
        document,
        "Le module Investissement de Najahni constitue un excellent cas d'etude de genie logiciel, car il combine une richesse fonctionnelle elevee et une architecture suffisamment propre pour etre analysee avec les outils du cours. La partie Scrum montre qu'il est possible de decomposer un module reel en User Stories credibles, priorisees et livreables par sprint. La partie Architecture montre que les concepts de MVC, de 3 couches et de 5 couches ne sont pas seulement theoriques: ils structurent concretement le code et rendent le projet evolutif."
    )
    add_paragraph(
        document,
        "Ce dossier prouve ainsi deux choses. D'abord, la maitrise des concepts academiques: backlog, sprint, DoD, MVC, separation des couches, maintenabilite, reutilisabilite. Ensuite, la capacite a les appliquer a un projet professionnel tangible. C'est precisement cette articulation entre theorie et implementation qui caracterise un travail de genie logiciel mature."
    )

    add_section_title(document, "Annexe - Exemple court d'interaction applicative", 1)
    add_paragraph(document, "Extrait conceptuel 1: calcul de risque", bold=True)
    add_paragraph(
        document,
        "InvestmentAdvancedController::riskCompute() recupere l'opportunite cible, interroge EconomicApiService pour les donnees pays, appelle EconomicRiskEngine::calculateFullRisk(), persiste le score sur InvestmentOpportunity puis renvoie un JSON contenant score, niveau, recommandation et facteurs. Cet enchainement illustre parfaitement l'orchestration de la logique applicative au-dessus d'un service metier specialiste.",
    )
    add_paragraph(document, "Extrait conceptuel 2: signature numerique", bold=True)
    add_paragraph(
        document,
        "InvestmentContractController::sign() controle l'acces, valide le nom et le canvas, puis delegue a ContractSignatureService::sign() la construction du hash SHA-256 base sur le digest du contrat, le role du signataire, la date, l'IP, le user-agent et l'image de signature. On voit ici un design propre: le controleur orchestre, le service metier calcule, l'entite conserve la preuve.",
    )


def generate_document():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document = Document()
    configure_styles(document)
    configure_page(document)
    add_title_page(document)

    backlog = build_product_backlog()

    add_intro_sections(document)
    add_scrum_part(document, backlog)
    add_architecture_part(document)
    add_conclusion(document)

    document.save(str(OUTPUT_FILE))
    print(f"DOCX generated: {OUTPUT_FILE}")


if __name__ == "__main__":
    generate_document()